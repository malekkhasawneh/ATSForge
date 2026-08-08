from __future__ import annotations

import io
import json
import os
import re
from copy import deepcopy
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from flask import Flask, jsonify, render_template, request, send_file, url_for
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, KeepTogether, Paragraph, SimpleDocTemplate, Spacer

from content_pages import ALL_PAGES, GUIDE_PAGES, NAV_GROUPS
from resume_tailor import (TailorError, ai_tailor, create_fallback_docx, extract_docx,
                           match_report, validate_docx)

app = Flask(__name__)

ADS_TXT_RECORD = "google.com, pub-2767727782899451, DIRECT, f08c47fec0942fa0\n"


def site_url() -> str:
    """Return the configured public origin without trusting request headers."""
    return os.getenv("SITE_URL", "https://atsforge.org").strip().rstrip("/")


def public_url(path: str = "/") -> str:
    return f"{site_url()}{path if path.startswith('/') else '/' + path}"


def configured_text(name: str, default: str = "") -> str:
    """Read a short display setting without allowing markup in templates."""
    return re.sub(r"[<>]", "", os.getenv(name, default)).strip()


def site_contact_details() -> dict[str, str]:
    """Return public support and operator details configured for the live site."""
    return {
        "support_email": configured_text("SUPPORT_EMAIL", "support@atsforge.org"),
        "legal_entity_name": configured_text("LEGAL_ENTITY_NAME", "ATSForge"),
        "legal_address": configured_text("LEGAL_ADDRESS"),
    }


def page_with_site_details(page_data: dict[str, Any]) -> dict[str, Any]:
    """Fill public-policy placeholders while keeping source page content reusable."""
    details = site_contact_details()
    replacements = {
        "{{SUPPORT_EMAIL}}": details["support_email"],
        "{{LEGAL_ENTITY_NAME}}": details["legal_entity_name"],
        "{{LEGAL_ADDRESS}}": details["legal_address"],
    }

    def replace(value: Any) -> Any:
        if isinstance(value, str):
            for token, replacement in replacements.items():
                value = value.replace(token, replacement)
            return value
        if isinstance(value, tuple):
            return tuple(replace(item) for item in value)
        if isinstance(value, list):
            return [replace(item) for item in value]
        if isinstance(value, dict):
            return {key: replace(item) for key, item in value.items()}
        return value

    return replace(deepcopy(page_data))

ACTION_VERBS = [
    "Accelerated", "Achieved", "Built", "Coordinated", "Delivered", "Designed",
    "Developed", "Directed", "Enhanced", "Implemented", "Improved", "Launched",
    "Led", "Optimized", "Reduced", "Resolved", "Scaled", "Streamlined",
]

RESUME_TEMPLATES = {
    "professional": {"name": "Professional", "font": "Arial", "pdf_font": "Helvetica", "color": "18221D", "name_size": 21, "align": "center"},
    "modern": {"name": "Modern", "font": "Arial", "pdf_font": "Helvetica", "color": "235F47", "name_size": 22, "align": "left"},
    "minimal": {"name": "Minimal", "font": "Arial", "pdf_font": "Helvetica", "color": "333333", "name_size": 19, "align": "left"},
    "executive": {"name": "Executive", "font": "Georgia", "pdf_font": "Times-Roman", "color": "20354A", "name_size": 23, "align": "center"},
}


def clean(value: Any, limit: int = 5000) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()[:limit]


def normalize_items(items: Any, fields: list[str], max_items: int = 10) -> list[dict[str, str]]:
    if not isinstance(items, list):
        return []
    result = []
    for item in items[:max_items]:
        if isinstance(item, dict):
            row = {field: clean(item.get(field), 1200) for field in fields}
            if any(row.values()):
                result.append(row)
    return result


def build_resume(payload: dict[str, Any]) -> dict[str, Any]:
    basics = payload.get("basics") if isinstance(payload.get("basics"), dict) else {}
    data = {
        "template": clean(payload.get("template"), 30) if clean(payload.get("template"), 30) in RESUME_TEMPLATES else "professional",
        "basics": {key: clean(basics.get(key), 300) for key in
                   ["name", "title", "email", "phone", "location", "linkedin", "website", "summary"]},
        "experience": normalize_items(payload.get("experience"), ["title", "company", "location", "start", "end", "highlights"]),
        "education": normalize_items(payload.get("education"), ["degree", "school", "location", "start", "end", "details"]),
        "projects": normalize_items(payload.get("projects"), ["name", "link", "description"]),
        "skills": [clean(x, 80) for x in payload.get("skills", [])[:40] if clean(x, 80)] if isinstance(payload.get("skills"), list) else [],
        "languages": [clean(x, 80) for x in payload.get("languages", [])[:12] if clean(x, 80)] if isinstance(payload.get("languages"), list) else [],
    }
    return data


def ats_analysis(data: dict[str, Any]) -> dict[str, Any]:
    basics = data["basics"]
    searchable = " ".join([
        basics.get("title", ""), basics.get("summary", ""), " ".join(data["skills"]),
        " ".join(data["languages"]), " ".join(x.get("title", "") for x in data["experience"]),
        " ".join(x.get("highlights", "") for x in data["experience"]),
        " ".join(x.get("degree", "") for x in data["education"]),
        " ".join(x.get("details", "") for x in data["education"]),
        " ".join(x.get("description", "") for x in data["projects"]),
    ])
    checks = {
        "Contact details": bool(basics.get("email") and basics.get("phone")),
        "Professional summary": len(basics.get("summary", "").split()) >= 20,
        "Work experience": bool(data["experience"]),
        "Measurable impact": bool(re.search(r"\b\d+(?:[.,]\d+)?%?|\$\d+", searchable)),
        "Relevant skills": len(data["skills"]) >= 5,
    }
    completeness = sum(checks.values()) / len(checks) * 80
    clarity = 20 if all(len(x.get("highlights", "").split()) <= 100 for x in data["experience"]) else 10
    score = min(100, round(completeness + clarity))
    return {"score": score, "checks": checks}


def set_bottom_border(paragraph, color: str = "555555", size: str = "6") -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def split_bullets(text: str) -> list[str]:
    return [part.strip() for part in re.split(r"\n|•", text) if part.strip()]


def create_docx(data: dict[str, Any]) -> io.BytesIO:
    doc = Document()
    template = RESUME_TEMPLATES.get(data.get("template", "professional"), RESUME_TEMPLATES["professional"])
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(.60)
    section.left_margin = section.right_margin = Inches(.70)
    styles = doc.styles
    styles["Normal"].font.name = template["font"]
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), template["font"])
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].paragraph_format.space_after = Pt(2.5)
    styles["Normal"].paragraph_format.line_spacing = 1.05
    bullet_style = styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet_style.base_style = styles["Normal"]
    bullet_style.paragraph_format.left_indent = Inches(.19)
    bullet_style.paragraph_format.first_line_indent = Inches(-.13)
    bullet_style.paragraph_format.space_after = Pt(1.5)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER if template["align"] == "center" else WD_ALIGN_PARAGRAPH.LEFT
    run = name.add_run(data["basics"]["name"] or "YOUR NAME")
    run.bold = True
    run.font.size = Pt(template["name_size"])
    from docx.shared import RGBColor
    run.font.color.rgb = RGBColor.from_string(template["color"])
    name.paragraph_format.space_after = Pt(1)
    if data["basics"].get("title"):
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER if template["align"] == "center" else WD_ALIGN_PARAGRAPH.LEFT
        title_run = title.add_run(data["basics"]["title"])
        title_run.bold = True
        title_run.font.size = Pt(10)
        title.paragraph_format.space_after = Pt(1)
    contact_values = [data["basics"].get(k) for k in ["location", "phone", "email", "linkedin", "website"]]
    contact = doc.add_paragraph(" | ".join(x for x in contact_values if x))
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER if template["align"] == "center" else WD_ALIGN_PARAGRAPH.LEFT
    contact.paragraph_format.space_after = Pt(8)
    if data.get("template") != "minimal":
        set_bottom_border(contact, template["color"], "8")

    def heading(label: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(label.upper())
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string(template["color"])
        p.style = styles["Normal"]
        set_bottom_border(p)

    if data["basics"].get("summary"):
        heading("Professional Summary")
        doc.add_paragraph(data["basics"]["summary"])
    if data["skills"]:
        heading("Core Skills")
        doc.add_paragraph(" • ".join(data["skills"]))
    if data["experience"]:
        heading("Professional Experience")
        for item in data["experience"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(item["title"])
            r.bold = True
            company_line = " — ".join(x for x in [item["company"], item["location"]] if x)
            if company_line:
                p.add_run(" | " + company_line)
            dates = " – ".join(x for x in [item["start"], item["end"]] if x)
            if dates:
                p.add_run(" | " + dates)
            for bullet in split_bullets(item["highlights"]):
                bp = doc.add_paragraph(style="Resume Bullet")
                bp.add_run("•  " + bullet)
    if data["projects"]:
        heading("Selected Projects")
        for item in data["projects"]:
            p = doc.add_paragraph()
            r = p.add_run(item["name"])
            r.bold = True
            if item["link"]:
                p.add_run(" | " + item["link"])
            if item["description"]:
                p.add_run(" — " + item["description"])
    if data["education"]:
        heading("Education")
        for item in data["education"]:
            title = " — ".join(x for x in [item["degree"], item["school"]] if x)
            dates = " – ".join(x for x in [item["start"], item["end"]] if x)
            p = doc.add_paragraph(title + ((" | " + dates) if dates else ""))
            if item["details"]:
                p.add_run(" | " + item["details"])
    if data["languages"]:
        heading("Languages")
        doc.add_paragraph(" • ".join(data["languages"]))
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def pdf_text(value: str) -> str:
    from xml.sax.saxutils import escape
    return escape(value or "")


def create_pdf(data: dict[str, Any]) -> io.BytesIO:
    out = io.BytesIO()
    template = RESUME_TEMPLATES.get(data.get("template", "professional"), RESUME_TEMPLATES["professional"])
    pdf_font = template["pdf_font"]
    pdf_bold = "Times-Bold" if pdf_font == "Times-Roman" else "Helvetica-Bold"
    pdf_align = TA_CENTER if template["align"] == "center" else 0
    accent = colors.HexColor("#" + template["color"])
    # Word uses US Letter by default. Using the identical page size and margins
    # keeps wrapping, section positions, and page breaks consistent across exports.
    document = SimpleDocTemplate(out, pagesize=letter, rightMargin=.70*inch, leftMargin=.70*inch,
        topMargin=.60*inch, bottomMargin=.60*inch,
        title=f"{data['basics'].get('name') or 'Candidate'} — Resume")
    base = getSampleStyleSheet()
    name_style = ParagraphStyle("ResumeName", parent=base["Normal"], fontName=pdf_bold, fontSize=template["name_size"], leading=template["name_size"]+2, alignment=pdf_align, textColor=accent, spaceAfter=1)
    center_style = ParagraphStyle("ResumeCenter", parent=base["Normal"], fontName=pdf_font, fontSize=9.5, leading=11, alignment=pdf_align, spaceAfter=1)
    title_style = ParagraphStyle("ResumeTitle", parent=center_style, fontName=pdf_bold, fontSize=10)
    heading_style = ParagraphStyle("ResumeHeading", parent=base["Normal"], fontName=pdf_bold, fontSize=10, textColor=accent, leading=12, spaceBefore=7, spaceAfter=1)
    body_style = ParagraphStyle("ResumeBody", parent=base["Normal"], fontName=pdf_font, fontSize=9.5, leading=11.5, spaceAfter=2.5)
    entry_style = ParagraphStyle("ResumeEntry", parent=body_style, fontName=pdf_bold, spaceAfter=1)
    bullet_style = ParagraphStyle("ResumeBullet", parent=body_style, leftIndent=12, firstLineIndent=-8, spaceAfter=1.5)
    story, basics = [], data["basics"]
    story.append(Paragraph(pdf_text(basics.get("name") or "YOUR NAME"), name_style))
    if basics.get("title"): story.append(Paragraph(pdf_text(basics["title"]), title_style))
    contacts = [basics.get(k) for k in ["location","phone","email","linkedin","website"] if basics.get(k)]
    if contacts: story.append(Paragraph(" &nbsp;|&nbsp; ".join(pdf_text(x) for x in contacts), center_style))
    story.append(Spacer(1, 3))
    if data.get("template") != "minimal":
        story.append(HRFlowable(width="100%", thickness=.8, color=accent, spaceAfter=5))
    def heading(label):
        story.append(Paragraph(pdf_text(label.upper()), heading_style))
        story.append(HRFlowable(width="100%", thickness=.55, color=accent, spaceAfter=3))
    if basics.get("summary"):
        heading("Professional Summary"); story.append(Paragraph(pdf_text(basics["summary"]), body_style))
    if data["skills"]:
        heading("Core Skills"); story.append(Paragraph(" • ".join(pdf_text(x) for x in data["skills"]), body_style))
    if data["experience"]:
        heading("Professional Experience")
        for item in data["experience"]:
            dates = " – ".join(x for x in [item["start"],item["end"]] if x)
            company_line = " — ".join(x for x in [item["company"],item["location"]] if x)
            line = f"<b>{pdf_text(item['title'])}</b>"
            if company_line: line += " | " + pdf_text(company_line)
            if dates: line += " | " + pdf_text(dates)
            block = [Paragraph(line, body_style)]
            block += [Paragraph("• &nbsp;"+pdf_text(b), bullet_style) for b in split_bullets(item["highlights"])]
            story.append(KeepTogether(block))
    if data["projects"]:
        heading("Selected Projects")
        for item in data["projects"]:
            label = " | ".join(pdf_text(x) for x in [item["name"],item["link"]] if x)
            desc = (" — "+pdf_text(item["description"])) if item["description"] else ""
            story.append(Paragraph(f"<b>{label}</b>{desc}", body_style))
    if data["education"]:
        heading("Education")
        for item in data["education"]:
            label = " — ".join(pdf_text(x) for x in [item["degree"],item["school"]] if x)
            dates = " – ".join(x for x in [item["start"],item["end"]] if x)
            line = label + ((" | "+pdf_text(dates)) if dates else "")
            if item["details"]: line += " | " + pdf_text(item["details"])
            story.append(Paragraph(line, body_style))
    if data["languages"]:
        heading("Languages"); story.append(Paragraph(" • ".join(pdf_text(x) for x in data["languages"]), body_style))
    document.build(story); out.seek(0); return out


@app.get("/")
def home():
    page_data = {
        "title": "Free ATS Resume Builder with DOCX and PDF Export",
        "description": "Build a clear, ATS-friendly resume, compare it with a job description, and export an editable DOCX or text-based PDF.",
        "faq": [],
    }
    schema = {
        "@context": "https://schema.org", "@graph": [
            {"@type": "WebSite", "name": "ATSForge", "url": public_url()},
            {"@type": "WebApplication", "name": "ATSForge", "applicationCategory": "BusinessApplication",
             "operatingSystem": "Web", "url": public_url(), "description": page_data["description"],
             "offers": {"@type": "Offer", "price": "0", "priceCurrency": "USD"}},
        ],
    }
    return render_template("index.html", page=page_data, schema=schema)


@app.get("/resume-tailor")
def resume_tailor_page():
    page_data = {
        "title": "AI Resume Tailor for ATS Job Matching",
        "description": "Upload a DOCX resume and tailor it to a job description with evidence-preserving AI and ATS-readable Word formatting.",
        "faq": [],
    }
    schema = {
        "@context": "https://schema.org", "@type": "WebApplication", "name": "ATSForge Resume Tailor",
        "applicationCategory": "BusinessApplication", "operatingSystem": "Web",
        "description": "Tailor an existing DOCX resume to a job description with ATS-readable formatting.",
        "offers": {"@type": "Offer", "price": "0", "priceCurrency": "USD"},
        "url": public_url("/resume-tailor"),
    }
    return render_template("resume_tailor.html", page=page_data, schema=schema,
                           ai_configured=bool(os.getenv("HF_TOKEN", "").strip()))


@app.get("/resume-templates")
def resume_template_selector():
    page_data = {
        "title": "Choose an ATS-Friendly Resume Template",
        "description": "Choose a professional, modern, minimal, or executive ATS resume template, then add your information and export DOCX or PDF.",
        "faq": [],
    }
    schema = {
        "@context": "https://schema.org", "@type": "CollectionPage", "name": page_data["title"],
        "description": page_data["description"], "mainEntity": {"@type": "ItemList", "numberOfItems": len(RESUME_TEMPLATES)},
        "url": public_url("/resume-templates"),
    }
    return render_template("template_selector.html", page=page_data, schema=schema, templates=RESUME_TEMPLATES)


@app.post("/api/tailor-resume")
def tailor_resume_upload():
    if request.content_length and request.content_length > 12 * 1024 * 1024:
        return jsonify({"error": "The request is too large. Each DOCX file must be under 5 MB."}), 413
    upload = request.files.get("resume")
    job_description = re.sub(r"\s+", " ", request.form.get("job_description", "")).strip()
    job_upload = request.files.get("job_file")
    if upload is None or not upload.filename:
        return jsonify({"error": "Choose a DOCX resume to continue."}), 400
    if job_upload is not None and job_upload.filename:
        job_content = job_upload.read(5 * 1024 * 1024 + 1)
        try:
            validate_docx(job_upload.filename, job_content)
            uploaded_description, _ = extract_docx(job_content)
            job_description = re.sub(r"\s+", " ", uploaded_description).strip()
        except TailorError as exc:
            return jsonify({"error": "Job description: " + str(exc)}), 422
    if len(job_description.split()) < 30:
        return jsonify({"error": "Paste or upload a complete job description of at least 30 words."}), 400
    if len(job_description) > 20000:
        return jsonify({"error": "The job description must be 20,000 characters or fewer."}), 400
    content = upload.read(5 * 1024 * 1024 + 1)
    try:
        validate_docx(upload.filename, content)
        source_text, blocks = extract_docx(content)
        before = match_report(source_text, job_description)
        ai_result = ai_tailor(source_text, job_description)
        if ai_result is not None:
            normalized = build_resume(ai_result)
            output = create_docx(normalized)
            output_text = json.dumps(normalized, ensure_ascii=False)
            after = match_report(output_text, job_description)
            mode = "ai"
        else:
            output = create_fallback_docx(blocks)
            after = before
            mode = "format-only"
    except TailorError as exc:
        return jsonify({"error": str(exc)}), 422
    safe_stem = re.sub(r"[^A-Za-z0-9_-]+", "_", upload.filename.rsplit(".", 1)[0]).strip("_") or "Resume"
    response = send_file(output, as_attachment=True, download_name=f"{safe_stem}_Tailored_ATS.docx",
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response.headers["X-Tailor-Mode"] = mode
    response.headers["X-ATS-Score-Before"] = str(before["score"])
    response.headers["X-ATS-Score-After"] = str(after["score"])
    response.headers["X-Job-Match-Before"] = str(before["score"])
    response.headers["X-Job-Match-After"] = str(after["score"])
    response.headers["X-Matched-Keywords"] = ", ".join(after["matched"][:10])
    response.headers["X-Missing-Keywords"] = ", ".join(after["missing"][:10])
    response.headers["X-Content-Type-Options"] = "nosniff"
    return response


@app.context_processor
def public_navigation():
    return {
        "all_pages": ALL_PAGES,
        "nav_groups": NAV_GROUPS,
        "site_url": site_url(),
        **site_contact_details(),
        "google_site_verification": os.getenv("GOOGLE_SITE_VERIFICATION", "").strip(),
    }


@app.get("/resources/<slug>")
def content_page(slug: str):
    page_data = ALL_PAGES.get(slug)
    if not page_data:
        return render_template("404.html"), 404
    page_data = page_with_site_details(page_data)
    words = " ".join([page_data["intro"]] + [f"{heading} {body}" for heading, body in page_data["sections"]]
                     + [f"{question} {answer}" for question, answer in page_data.get("faq", [])])
    word_count = len(re.findall(r"\b[\w’'-]+\b", words))
    page_url = public_url(f"/resources/{slug}")
    schema = {
        "@context": "https://schema.org", "@graph": [
            {"@type": "Article", "headline": page_data["title"], "description": page_data["description"],
             "dateModified": "2026-08-08", "datePublished": "2026-08-06",
             "author": {"@type": "Organization", "name": "ATSForge"},
             "publisher": {"@type": "Organization", "name": "ATSForge"}, "mainEntityOfPage": page_url},
            {"@type": "BreadcrumbList", "itemListElement": [
                {"@type": "ListItem", "position": 1, "name": "Home", "item": public_url()},
                {"@type": "ListItem", "position": 2, "name": page_data["category"], "item": page_url},
                {"@type": "ListItem", "position": 3, "name": page_data["title"], "item": page_url},
            ]},
        ],
    }
    faq_schema = None
    if page_data.get("faq"):
        faq_schema = {"@context": "https://schema.org", "@type": "FAQPage", "mainEntity": [
            {"@type": "Question", "name": q, "acceptedAnswer": {"@type": "Answer", "text": a}}
            for q, a in page_data["faq"]
        ]}
    return render_template("content_page.html", page=page_data, schema=schema,
                           faq_schema=faq_schema, word_count=word_count)


@app.get("/sitemap.xml")
def sitemap():
    urls = [public_url(), public_url("/resume-tailor"), public_url("/resume-templates")] + [
        public_url(f"/resources/{slug}") for slug in ALL_PAGES
    ]
    body = "<?xml version=\"1.0\" encoding=\"UTF-8\"?>\n<urlset xmlns=\"http://www.sitemaps.org/schemas/sitemap/0.9\">" + "".join(
        f"<url><loc>{url}</loc><lastmod>2026-08-08</lastmod></url>" for url in urls) + "</urlset>"
    return app.response_class(body, mimetype="application/xml")


@app.get("/robots.txt")
def robots():
    body = "\n".join([
        "User-agent: *", "Allow: /", "Disallow: /api/", "Disallow: /health",
        f"Sitemap: {public_url('/sitemap.xml')}", "",
    ])
    return app.response_class(body, mimetype="text/plain")


@app.get("/ads.txt")
def ads_txt():
    """Declare ATSForge's authorized Google AdSense seller record."""
    return app.response_class(ADS_TXT_RECORD, mimetype="text/plain")


@app.errorhandler(404)
def not_found(_error):
    return render_template("404.html"), 404


@app.after_request
def set_response_headers(response):
    if request.path.startswith("/static/"):
        response.headers["Cache-Control"] = "public, max-age=604800"
    elif request.path.startswith("/api/"):
        # Resume contents and generated documents are sensitive and should not be cached.
        response.headers["Cache-Control"] = "no-store, private"
        response.headers["Pragma"] = "no-cache"
    response.headers.setdefault("X-Content-Type-Options", "nosniff")
    response.headers.setdefault("X-Frame-Options", "SAMEORIGIN")
    response.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
    response.headers.setdefault("Permissions-Policy", "camera=(), geolocation=(), microphone=()")
    return response


@app.post("/api/analyze")
def analyze():
    data = build_resume(request.get_json(silent=True) or {})
    return jsonify({"resume": data, "analysis": ats_analysis(data)})


@app.post("/api/download/docx")
def download_docx():
    data = build_resume(request.get_json(silent=True) or {})
    safe_name = re.sub(r"[^A-Za-z0-9_-]+", "_", data["basics"].get("name") or "resume").strip("_")
    return send_file(create_docx(data), as_attachment=True, download_name=f"{safe_name}_ATS_Resume.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.post("/api/download/pdf")
def download_pdf():
    data = build_resume(request.get_json(silent=True) or {})
    safe_name = re.sub(r"[^A-Za-z0-9_-]+", "_", data["basics"].get("name") or "resume").strip("_")
    return send_file(create_pdf(data), as_attachment=True, download_name=f"{safe_name}_ATS_Resume.pdf", mimetype="application/pdf")


@app.get("/health")
def health():
    return {"status": "ok", "year": datetime.now().year}


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(
        host="0.0.0.0",
        port=port,
        debug=False,
    )
