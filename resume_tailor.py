"""DOCX-only resume tailoring helpers, isolated from the existing builder."""

from __future__ import annotations

import io
import json
import os
import re
import zipfile
from collections import Counter
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

MAX_DOCX_BYTES = 5 * 1024 * 1024
MAX_UNCOMPRESSED_BYTES = 25 * 1024 * 1024
SECTION_NAMES = {
    "summary": "Professional Summary", "professional summary": "Professional Summary",
    "profile": "Professional Summary", "objective": "Professional Summary",
    "skills": "Core Skills", "core skills": "Core Skills", "technical skills": "Core Skills",
    "experience": "Professional Experience", "work experience": "Professional Experience",
    "professional experience": "Professional Experience", "employment": "Professional Experience",
    "education": "Education", "projects": "Selected Projects", "selected projects": "Selected Projects",
    "certifications": "Certifications", "languages": "Languages", "awards": "Awards",
}
STOP_WORDS = {
    "and", "the", "for", "with", "that", "this", "from", "your", "you", "our", "are", "will",
    "have", "has", "job", "role", "team", "work", "years", "using", "required", "preferred", "into",
    "who", "but", "not", "all", "can", "their", "skills", "experience", "about", "they", "was", "were",
    "hiring", "candidate", "candidates", "create", "improve", "lead", "maintain", "test", "looking", "join",
    "responsibilities", "requirements", "responsible", "include", "including", "ability", "strong", "excellent",
    "need", "needs", "seeking", "position", "opportunity", "qualification", "qualifications", "includes",
    "building", "build", "built", "must", "should",
}


class TailorError(ValueError):
    pass


def validate_docx(filename: str, content: bytes) -> None:
    if not filename.lower().endswith(".docx"):
        raise TailorError("Only .docx resume files are accepted.")
    if not content:
        raise TailorError("The uploaded resume is empty.")
    if len(content) > MAX_DOCX_BYTES:
        raise TailorError("The DOCX file must be 5 MB or smaller.")
    if not zipfile.is_zipfile(io.BytesIO(content)):
        raise TailorError("The file is not a valid DOCX document.")
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        total = sum(item.file_size for item in archive.infolist())
        if total > MAX_UNCOMPRESSED_BYTES or len(archive.infolist()) > 1500:
            raise TailorError("The DOCX package is too large or complex to process safely.")
        if "word/document.xml" not in archive.namelist():
            raise TailorError("The DOCX does not contain a readable Word document.")


def extract_docx(content: bytes) -> tuple[str, list[dict[str, Any]]]:
    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise TailorError("The DOCX could not be read. Open and resave it in Word, then try again.") from exc
    blocks: list[dict[str, Any]] = []
    seen = set()
    for paragraph in document.paragraphs:
        text = re.sub(r"\s+", " ", paragraph.text).strip()
        if text:
            style = (paragraph.style.name if paragraph.style else "") or ""
            blocks.append({"text": text, "style": style, "bullet": "list" in style.lower()})
            seen.add(text)
    # Many older resumes use tables. Read their text so no experience disappears.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = re.sub(r"\s+", " ", paragraph.text).strip()
                    if text and text not in seen:
                        blocks.append({"text": text, "style": "Table text", "bullet": False})
                        seen.add(text)
    text = "\n".join(block["text"] for block in blocks)
    if len(text.split()) < 20:
        raise TailorError("The resume contains too little selectable text. Scanned or image-only resumes are not supported.")
    return text[:30000], blocks[:600]


def keywords(text: str, limit: int = 24) -> list[str]:
    tokens = [token.rstrip("./-") for token in re.findall(r"[A-Za-z][A-Za-z0-9+#./-]{2,30}", text.lower())]
    counts = Counter(token for token in tokens if token not in STOP_WORDS)
    return [token for token, _ in counts.most_common(limit)]


def match_report(resume_text: str, job_description: str) -> dict[str, Any]:
    """Measure selected job-language coverage; it is not an employer ATS score."""
    job_keywords = keywords(job_description)
    source = resume_text.lower()
    matched = [word for word in job_keywords if re.search(rf"(?<!\w){re.escape(word)}(?!\w)", source)]
    missing = [word for word in job_keywords if word not in matched]
    score = round(100 * len(matched) / len(job_keywords)) if job_keywords else 0
    return {"score": score, "matched": matched, "missing": missing[:12], "keywords": job_keywords}


def safe_json(raw: str) -> dict[str, Any]:
    raw = raw.strip()
    raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw, flags=re.I)
    start, end = raw.find("{"), raw.rfind("}")
    if start < 0 or end <= start:
        raise TailorError("The AI provider returned an invalid response. Please try again.")
    try:
        value = json.loads(raw[start:end + 1])
    except json.JSONDecodeError as exc:
        raise TailorError("The AI provider returned malformed resume data. Please try again.") from exc
    if not isinstance(value, dict):
        raise TailorError("The AI response was not a resume object.")
    return value


def validate_grounding(source: str, tailored: dict[str, Any]) -> None:
    output = json.dumps(tailored, ensure_ascii=False)
    source_numbers = set(re.findall(r"(?<!\w)[£€$]?\d[\d,.:/%+-]*", source))
    output_numbers = set(re.findall(r"(?<!\w)[£€$]?\d[\d,.:/%+-]*", output))
    invented = sorted(output_numbers - source_numbers)
    if invented:
        raise TailorError("AI safety check rejected newly invented numbers: " + ", ".join(invented[:6]))
    source_folded = re.sub(r"\s+", " ", source).casefold()
    unsupported = []
    basics = tailored.get("basics", {}) if isinstance(tailored.get("basics"), dict) else {}
    if not str(basics.get("name", "")).strip():
        raise TailorError("AI safety check rejected a response with no candidate name.")
    for field in ("name", "email", "phone", "linkedin", "website"):
        value = str(basics.get(field, "")).strip()
        if value and re.sub(r"\s+", " ", value).casefold() not in source_folded:
            unsupported.append(f"{field}: {value}")
    for skill in tailored.get("skills", []) if isinstance(tailored.get("skills"), list) else []:
        value = str(skill).strip()
        if value and value.casefold() not in source_folded:
            unsupported.append(f"skill: {value}")
    identity_fields = {
        "experience": ("title", "company", "start", "end"),
        "education": ("degree", "school", "start", "end"),
        "projects": ("name",),
    }
    for section, fields in identity_fields.items():
        items = tailored.get(section, [])
        if not isinstance(items, list):
            continue
        for item in items:
            if not isinstance(item, dict):
                continue
            for field in fields:
                value = str(item.get(field, "")).strip()
                if value and re.sub(r"\s+", " ", value).casefold() not in source_folded:
                    unsupported.append(f"{section}.{field}: {value}")
    if unsupported:
        raise TailorError("AI safety check rejected unsupported identity or skill data: " + "; ".join(unsupported[:5]))


def ai_tailor(resume_text: str, job_description: str) -> dict[str, Any] | None:
    token = os.getenv("HF_TOKEN", "").strip()
    if not token:
        return None
    from huggingface_hub import InferenceClient
    model = os.getenv("HF_MODEL", "Qwen/Qwen2.5-7B-Instruct-1M")
    client = InferenceClient(provider="auto", api_key=token, timeout=90)
    system = """You are a senior resume editor specializing in ATS-readable resumes. Return JSON only.
NON-NEGOTIABLE EVIDENCE RULES:
- Never invent or infer an employer, title, date, degree, certification, tool, responsibility, number, metric, result, or skill.
- Use only facts explicitly present in SOURCE RESUME. The job description is context, never evidence.
- Preserve identity, contact information, employers, titles, education, dates, and all factual meaning.
- You may reorder supported evidence, remove irrelevant repetition, improve grammar, and mirror job terminology only where the source proves the same capability.
- Retain source terms that truthfully match the target job; do not reduce relevant keyword coverage merely to shorten the resume.
- Do not add a missing keyword merely to improve a score.
- Keep bullets concise, varied, specific, and human. Do not use first-person pronouns.
- Output a single-column resume object using exactly the requested keys."""
    prompt = f"""Tailor the source resume for the target job while obeying every evidence rule.

SOURCE RESUME:
{resume_text}

TARGET JOB DESCRIPTION:
{job_description[:15000]}

Return this JSON shape:
{{"basics":{{"name":"","title":"","email":"","phone":"","location":"","linkedin":"","website":"","summary":""}},"skills":[""],"experience":[{{"title":"","company":"","location":"","start":"","end":"","highlights":"one bullet per line"}}],"education":[{{"degree":"","school":"","location":"","start":"","end":"","details":""}}],"projects":[{{"name":"","link":"","description":""}}],"languages":[""],"change_log":[""]}}"""
    try:
        response = client.chat_completion(
            model=model, messages=[{"role": "system", "content": system}, {"role": "user", "content": prompt}],
            max_tokens=3500, temperature=0.2, top_p=0.85,
        )
        result = safe_json(response.choices[0].message.content or "")
        validate_grounding(resume_text, result)
        return result
    except TailorError:
        raise
    except Exception as exc:
        raise TailorError(f"Hugging Face tailoring is temporarily unavailable: {type(exc).__name__}.") from exc


def fallback_blocks(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Infer only presentation roles; never modify source wording."""
    result = []
    for index, block in enumerate(blocks):
        text = block["text"]
        normalized = re.sub(r"[^a-z ]", "", text.lower()).strip()
        heading = SECTION_NAMES.get(normalized)
        kind = "heading" if heading else "body"
        if block["bullet"] or re.match(r"^[•●▪◦-]\s+", text):
            kind = "bullet"
            text = re.sub(r"^[•●▪◦-]\s+", "", text)
        elif index == 0:
            kind = "name"
        elif index < 4 and ("@" in text or re.search(r"\+?\d[\d ()-]{7,}", text)):
            kind = "contact"
        elif heading:
            text = heading
        result.append({"text": text, "kind": kind})
    return result


def _bottom_border(paragraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "3"); bottom.set(qn("w:color"), "555555")
    borders.append(bottom); properties.append(borders)


def create_fallback_docx(blocks: list[dict[str, Any]]) -> io.BytesIO:
    """Reformat source wording into a clean, single-column ATS document."""
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(.60)
    section.left_margin = section.right_margin = Inches(.70)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"; normal.font.size = Pt(9.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.paragraph_format.space_after = Pt(2.5); normal.paragraph_format.line_spacing = 1.05
    bullet = document.styles.add_style("ATS Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet.base_style = normal; bullet.paragraph_format.left_indent = Inches(.19)
    bullet.paragraph_format.first_line_indent = Inches(-.13); bullet.paragraph_format.space_after = Pt(1.5)
    for block in fallback_blocks(blocks):
        kind, text = block["kind"], block["text"]
        paragraph = document.add_paragraph()
        if kind == "name":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(text); run.bold = True; run.font.size = Pt(21)
            paragraph.paragraph_format.space_after = Pt(1)
        elif kind == "contact":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER; paragraph.add_run(text)
        elif kind == "heading":
            paragraph.paragraph_format.space_before = Pt(7); paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(text.upper()); run.bold = True; run.font.size = Pt(10)
            _bottom_border(paragraph)
        elif kind == "bullet":
            paragraph.style = bullet; paragraph.add_run("•  " + text)
        else:
            paragraph.add_run(text)
    output = io.BytesIO(); document.save(output); output.seek(0); return output
